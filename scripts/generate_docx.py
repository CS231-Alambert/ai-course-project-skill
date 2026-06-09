#!/usr/bin/env python3
"""Generate a formatted .docx from a Markdown thesis/course-report.

Supports:
  - Hierarchical headings (# → ## → ### → ####)
  - Tables (GFM pipe syntax)
  - LaTeX formulas ($$ ... $$ blocks)
  - Image placeholders → auto-insert from image-map.json
  - Mermaid diagrams → rendered via mermaid-cli or placeholder
  - Code blocks → monospace Consolas
  - Chinese thesis defaults (宋体 body, 黑体 headings)

Usage:
  python generate_docx.py thesis.md output.docx [--image-map image-map.json]
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ── regex patterns ──────────────────────────────────────────
SCREENSHOT_RE = re.compile(r"^\[此处插入截图：(.+?)\]$")
# Also match: ![label](path) → extract label for image-map lookup
FIGURE_MD_RE = re.compile(r"^!\[(.+?)\]\((.+?)\)$")

# ── inline markdown cleanup ─────────────────────────────────
_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_AUTOLINK_RE = re.compile(r"<(https?://[^>]+)>")


def cleanup_text(text: str) -> str:
    """Strip inline markdown tokens, keep content."""
    c = text.replace("`", "")
    c = _LINK_RE.sub(r"\1", c)
    c = _AUTOLINK_RE.sub(r"\1", c)
    for pat, repl in [(r"\*\*(.+?)\*\*", r"\1"), (r"__(.+?)__", r"\1"),
                       (r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"\1"),
                       (r"(?<!_)_(?!\s)(.+?)(?<!\s)_(?!_)", r"\1")]:
        c = re.sub(pat, repl, c)
    return re.sub(r"\s+", " ", c).strip()


# ── style profile ───────────────────────────────────────────

def build_style_profile() -> dict[str, Any]:
    """AI课程论文格式 — 宋体标题/正文，1.5倍行距，匹配学校要求。"""

    def p_style(ea: str, la: str, sz: float, *, bold: bool = False,
                align: str = "left", ls: float = 1.5, fi: float = 0,
                li: float = 0, sb: float = 0, sa: float = 0,
                pb: bool = False) -> dict:
        return {
            "east_asia_font": ea, "latin_font": la, "size_pt": sz,
            "bold": bold, "alignment": align, "line_spacing": ls,
            "line_spacing_rule": "multiple" if ls != 1 else "single",
            "first_line_indent_pt": fi, "left_indent_pt": li,
            "space_before_pt": sb, "space_after_pt": sa,
            "page_break_before": pb,
        }

    def ctr(ea: str, la: str, sz: float, **kw) -> dict:
        return p_style(ea, la, sz, align="center", fi=0, **kw)

    return {"styles": {
        # 封面大标题 → 宋体二号加粗居中
        "title": ctr("宋体", "Times New Roman", 22, bold=True, sb=20, sa=20),
        # 一级标题 → 宋体三号 (16pt) 加粗
        "heading1": p_style("宋体", "Times New Roman", 16, bold=True, sb=12, sa=6),
        # 二级标题 → 宋体四号 (14pt) 加粗
        "heading2": p_style("宋体", "Times New Roman", 14, bold=True, sb=10, sa=4),
        # 三级标题 → 宋体小四 (12pt) 加粗
        "heading3": p_style("宋体", "Times New Roman", 12, bold=True, sb=8, sa=4),
        # 摘要/目录/参考文献/致谢/附录 → 宋体三号居中加粗
        "abstract_heading_cn": ctr("宋体", "Times New Roman", 16, bold=True, pb=True),
        "abstract_heading_en": ctr("Times New Roman", "Times New Roman", 16, bold=True, pb=True),
        "toc_heading": ctr("宋体", "Times New Roman", 16, bold=True, pb=True),
        "references_heading": ctr("宋体", "Times New Roman", 16, bold=True, pb=True),
        "ack_heading": ctr("宋体", "Times New Roman", 16, bold=True, pb=True),
        "appendix_heading": ctr("宋体", "Times New Roman", 16, bold=True, pb=True),
        # 正文 → 宋体小四 (12pt)，首行缩进2字符(24pt)，1.5倍行距
        "body_cn": p_style("宋体", "Times New Roman", 12, fi=24),
        "body_en": p_style("Times New Roman", "Times New Roman", 12, fi=21),
        # 特殊元素
        "keywords_cn_label": {"east_asia_font": "宋体", "latin_font": "Times New Roman", "size_pt": 12, "bold": True},
        "keywords_cn_content": {"east_asia_font": "宋体", "latin_font": "Times New Roman", "size_pt": 12},
        "keywords_en_label": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12, "bold": True},
        "keywords_en_content": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12},
        "keywords_paragraph": p_style("宋体", "Times New Roman", 10.5, fi=0),
        # 图表标题 → 宋体五号 (10.5pt)，单倍行距
        "figure_caption": ctr("宋体", "Times New Roman", 10.5, ls=1),
        "table_caption": ctr("宋体", "Times New Roman", 10.5, ls=1),
        "table_text": ctr("宋体", "Times New Roman", 10.5),
        # 参考文献 → 宋体五号，悬挂缩进
        "references_body": p_style("宋体", "Times New Roman", 10.5, fi=-21, li=21),
        # 公式
        "equation": ctr("Times New Roman", "Times New Roman", 12, ls=1.5),
        # 缺失素材占位
        "missing_asset": ctr("楷体", "Times New Roman", 10.5),
        # 代码块
        "code": p_style("Times New Roman", "Times New Roman", 10.5, fi=0),
    }}


SPECIAL_HEADINGS = {
    "摘要": "abstract_heading_cn", "abstract": "abstract_heading_en",
    "目录": "toc_heading", "参考文献": "references_heading",
    "致谢": "ack_heading", "附录": "appendix_heading",
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ── OOXML helpers ───────────────────────────────────────────

def set_run_font(run, ea: str, la: str, sz: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(sz)
    run.font.name = la
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:ascii"), la)
    rFonts.set(qn("w:hAnsi"), la)


def apply_para_style(para, s: dict) -> None:
    para.alignment = ALIGN_MAP.get(s.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE if s.get("line_spacing_rule") == "single" else WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = s.get("line_spacing", 1.25)
    pf.first_line_indent = Pt(s.get("first_line_indent_pt", 0))
    pf.left_indent = Pt(s.get("left_indent_pt", 0))
    pf.space_before = Pt(s.get("space_before_pt", 0))
    pf.space_after = Pt(s.get("space_after_pt", 0))
    if s.get("page_break_before"):
        pPr = para._p.get_or_add_pPr()
        pb = OxmlElement("w:pageBreakBefore")
        pPr.append(pb)
    for run in para.runs:
        set_run_font(run, s["east_asia_font"], s["latin_font"], s["size_pt"], bold=s.get("bold", False))


def add_page_setup(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


# ── element builders ────────────────────────────────────────

def add_keywords(para, label: str, content: str, S: dict) -> None:
    apply_para_style(para, S["keywords_paragraph"])
    is_cn = label.startswith("关键词")
    ls = S["keywords_cn_label" if is_cn else "keywords_en_label"]
    cs = S["keywords_cn_content" if is_cn else "keywords_en_content"]
    r = para.add_run(label)
    set_run_font(r, ls["east_asia_font"], ls["latin_font"], ls["size_pt"], bold=ls.get("bold", False))
    if content:
        sep = "" if label.endswith(("：", ":")) else " "
        r2 = para.add_run(f"{sep}{cleanup_text(content)}")
        set_run_font(r2, cs["east_asia_font"], cs["latin_font"], cs["size_pt"])


def add_md_table(doc: Document, lines: list[str], S: dict) -> None:
    rows = []
    for ln in lines:
        cells = [cleanup_text(c.strip()) for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[:\- ]+", c or "") for c in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return
    hdr, body = rows[0], rows[1:]
    tbl = doc.add_table(rows=1 + len(body), cols=len(hdr))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, t in enumerate(hdr):
        tbl.rows[0].cells[i].text = t
    for ri, row in enumerate(body, 1):
        for ci, t in enumerate(row):
            if ci < len(tbl.rows[ri].cells):
                tbl.rows[ri].cells[ci].text = t
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                apply_para_style(p, S["table_text"])


def add_image(doc: Document, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(14))


def add_missing(doc: Document, label: str, S: dict) -> None:
    p = doc.add_paragraph()
    p.add_run(f"【待补：{label}】")
    apply_para_style(p, S["missing_asset"])


# ── main builder ────────────────────────────────────────────

def build_docx(source: Path, image_map: dict[str, Path],
               formula_mode: str = "latex_text") -> Document:
    S = build_style_profile()["styles"]
    doc = Document()
    add_page_setup(doc.sections[0])

    lines = source.read_text(encoding="utf-8").splitlines()
    cur_section = ""        # current ## heading, for context
    seen_content = False    # has any real content been added?
    in_code = False
    code_lang = ""
    pending_mermaid = False
    i = 0

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # ── code fence ──────────────────────────────────
        if in_code:
            if s.startswith("```"):
                in_code = False
                if code_lang == "mermaid":
                    pending_mermaid = True
                code_lang = ""
            elif code_lang != "mermaid":
                p = doc.add_paragraph()
                p.add_run(line.rstrip())
                apply_para_style(p, S["code"])
            i += 1
            continue

        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue

        if s.startswith("```"):
            in_code = True
            code_lang = s[3:].strip().lower()
            i += 1
            continue

        # ── formula block $$ ... $$ ──────────────────────
        if s == "$$":
            fb = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                fb.append(lines[i].rstrip())
                i += 1
            if i < len(lines) and lines[i].strip() == "$$":
                i += 1
            ft = "\n".join(fb).strip()
            if ft:
                p = doc.add_paragraph()
                p.add_run(ft)
                apply_para_style(p, S["equation"])
                seen_content = True
            continue

        # inline formula $$...$$
        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            p = doc.add_paragraph()
            p.add_run(s[2:-2].strip())
            apply_para_style(p, S["equation"])
            seen_content = True
            i += 1
            continue

        # ── headings ────────────────────────────────────
        if s.startswith("# "):
            p = doc.add_paragraph()
            p.add_run(cleanup_text(s[2:]))
            apply_para_style(p, S["title"])
            seen_content = True
            i += 1
            continue

        if s.startswith("## "):
            text = cleanup_text(s[3:])
            key = SPECIAL_HEADINGS.get(text.replace(" ", "").lower(), "heading1")
            p = doc.add_paragraph()
            p.add_run(text)
            style = S[key]
            if seen_content and key == "heading1":
                style = {**S["heading1"], "page_break_before": True}
            apply_para_style(p, style)
            cur_section = text
            seen_content = True
            i += 1
            continue

        if s.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(cleanup_text(s[4:]))
            apply_para_style(p, S["heading2"])
            seen_content = True
            i += 1
            continue

        if s.startswith("#### "):
            p = doc.add_paragraph()
            p.add_run(cleanup_text(s[5:]))
            apply_para_style(p, S["heading3"])
            seen_content = True
            i += 1
            continue

        # ── keywords ────────────────────────────────────
        kw = re.match(r"^(关键词[:：]|Keywords[:：])\s*(.*)$", s)
        if kw:
            p = doc.add_paragraph()
            add_keywords(p, kw.group(1), kw.group(2), S)
            seen_content = True
            i += 1
            continue

        # ── images ──────────────────────────────────────
        sm = SCREENSHOT_RE.match(s)
        if sm:
            label = sm.group(1).strip()
            ip = image_map.get(label)
            if ip and ip.exists():
                add_image(doc, ip)
            else:
                add_missing(doc, label, S)
            seen_content = True
            i += 1
            continue

        fm = FIGURE_MD_RE.match(s)
        if fm:
            label = fm.group(1).strip()
            ip = image_map.get(label)
            if ip and ip.exists():
                add_image(doc, ip)
            else:
                add_missing(doc, label, S)
            seen_content = True
            i += 1
            continue

        # ── figure / table captions ─────────────────────
        if re.match(r"^图\s*\d+(\.\d+)?", s):
            cap = cleanup_text(s)
            if pending_mermaid:
                ip = image_map.get(cap)
                if ip and ip.exists():
                    add_image(doc, ip)
                else:
                    add_missing(doc, cap, S)
                pending_mermaid = False
            p = doc.add_paragraph()
            p.add_run(cap)
            apply_para_style(p, S["figure_caption"])
            seen_content = True
            i += 1
            continue

        if re.match(r"^表\s*\d+(\.\d+)?", s):
            p = doc.add_paragraph()
            p.add_run(cleanup_text(s))
            apply_para_style(p, S["table_caption"])
            seen_content = True
            i += 1
            continue

        # ── GFM table ───────────────────────────────────
        if s.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            add_md_table(doc, tbl_lines, S)
            seen_content = True
            continue

        # ── body paragraph ──────────────────────────────
        p = doc.add_paragraph()
        p.add_run(cleanup_text(s))
        ref_sec = cur_section.replace(" ", "") == "参考文献"
        if ref_sec and re.match(r"^\[\d+\]", s):
            apply_para_style(p, S["references_body"])
        elif cur_section.lower() == "abstract":
            apply_para_style(p, S["body_en"])
        else:
            apply_para_style(p, S["body_cn"])
        seen_content = True
        i += 1

    return doc


# ── CLI ─────────────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser(description="Generate formatted .docx from Markdown thesis/report.")
    ap.add_argument("source", type=Path, help="Input .md file")
    ap.add_argument("output", type=Path, help="Output .docx file")
    ap.add_argument("--image-map", type=Path, default=None,
                    help="JSON file mapping image labels → file paths")
    ap.add_argument("--formula-mode", choices=["latex_text", "formula_image"],
                    default="latex_text", help="Formula rendering mode")
    args = ap.parse_args()

    imap: dict[str, Path] = {}
    if args.image_map and args.image_map.exists():
        raw = json.loads(args.image_map.read_text(encoding="utf-8"))
        imap = {k: Path(v) for k, v in raw.items()}

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc = build_docx(args.source, imap, formula_mode=args.formula_mode)
    doc.save(str(args.output))
    print(f"✓  {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
