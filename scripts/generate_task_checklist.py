#!/usr/bin/env python3
"""Generate a task checklist DOCX for the course project paper.

Reads the thesis markdown, finds all missing assets (screenshots, figures),
and produces a formatted DOCX with:
  1. Screenshot tasks — what to capture, from which source file
  2. File delivery instructions — where to place each file
  3. Final verification checklist

Output to the user's desktop as: <论文标题>-待办清单.docx
"""

from __future__ import annotations

import argparse
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def find_desktop() -> Path:
    """Try common desktop paths on Windows/WSL."""
    candidates = []
    # Scan /mnt/c/Users for actual user desktops
    try:
        for d in Path("/mnt/c/Users").iterdir():
            if d.name in ("Public", "Default", "Default User", "All Users"):
                continue
            desk = d / "Desktop"
            if desk.is_dir():
                candidates.append(desk)
    except Exception:
        pass
    candidates += [
        Path(os.environ.get("USERPROFILE", "")) / "Desktop",
        Path("/mnt/c/Users/Public/Desktop"),
        Path.home() / "Desktop",
    ]
    for c in candidates:
        if c.is_dir():
            return c
    return Path.home() / "Desktop"


# ── style helpers (same as generate_docx.py) ──────────────────

def p_style(ea: str, la: str, sz: float, *, bold: bool = False,
            align: str = "left", ls: float = 1.5, fi: float = 0,
            sb: float = 0, sa: float = 0) -> dict:
    return {
        "east_asia_font": ea, "latin_font": la, "size_pt": sz,
        "bold": bold, "alignment": align, "line_spacing": ls,
        "line_spacing_rule": "multiple" if ls != 1 else "single",
        "first_line_indent_pt": fi, "left_indent_pt": 0,
        "space_before_pt": sb, "space_after_pt": sa,
        "page_break_before": False,
    }


def ctr(ea: str, la: str, sz: float, **kw) -> dict:
    return p_style(ea, la, sz, align="center", fi=0, **kw)


STYLES = {
    "title": ctr("黑体", "Times New Roman", 22, bold=True, sb=20, sa=20),
    "heading1": p_style("宋体", "Times New Roman", 16, bold=True, sb=12, sa=6),
    "heading2": p_style("宋体", "Times New Roman", 14, bold=True, sb=8, sa=4),
    "body": p_style("宋体", "Times New Roman", 12, fi=24),
    "code_path": p_style("Times New Roman", "Times New Roman", 10.5, fi=0, ls=1),
    "todo": p_style("宋体", "Times New Roman", 12, fi=0, sb=2, sa=2),
    "note": p_style("楷体", "Times New Roman", 10.5, fi=0),
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
}


def set_run_font(run, ea: str, la: str, sz: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(sz)
    run.font.name = la
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:ascii"), la)
    rFonts.set(qn("w:hAnsi"), la)


def apply_style(para, s: dict) -> None:
    para.alignment = ALIGN_MAP.get(s.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE if s.get("line_spacing_rule") == "single" else WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = s.get("line_spacing", 1.5)
    pf.first_line_indent = Pt(s.get("first_line_indent_pt", 0))
    pf.space_before = Pt(s.get("space_before_pt", 0))
    pf.space_after = Pt(s.get("space_after_pt", 0))
    for run in para.runs:
        set_run_font(run, s["east_asia_font"], s["latin_font"], s["size_pt"],
                     bold=s.get("bold", False))


def add_heading(doc: Document, text: str, style_key: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    apply_style(p, STYLES[style_key])


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    apply_style(p, STYLES["body"])


def add_path(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    apply_style(p, STYLES["code_path"])


def add_checkbox(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph()
    prefix = "    " if indent else ""
    p.add_run(f"{prefix}☐ {text}")
    apply_style(p, STYLES["todo"])


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"💡 {text}")
    apply_style(p, STYLES["note"])


# ── parser ───────────────────────────────────────────────────

SCREENSHOT_RE = re.compile(r"\[此处插入截图：(.+?)\]")
FIGURE_RE = re.compile(r"!\[(.+?)\]\((.+?)\)")


def parse_tasks(md_path: Path) -> dict[str, Any]:
    """Parse markdown to extract all tasks and assets."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    screenshots = []
    figures = []
    ref_count = 0
    paper_title = md_path.stem

    for i, line in enumerate(lines):
        s = line.strip()

        # Screenshot placeholders
        m = SCREENSHOT_RE.search(s)
        if m:
            # Find preceding heading for context
            context = ""
            for j in range(i - 1, max(i - 10, -1), -1):
                if lines[j].strip().startswith("### "):
                    context = lines[j].strip()[4:]
                    break
                elif lines[j].strip().startswith("## "):
                    context = lines[j].strip()[3:]
                    break
            screenshots.append({
                "label": m.group(1),
                "context": context,
                "line": i + 1,
            })

        # Existing figures (real images)
        m = FIGURE_RE.match(s)
        if m:
            figures.append({
                "label": m.group(1),
                "path": m.group(2),
                "line": i + 1,
            })

        # Count references
        if re.match(r"^\[\d+\]", s):
            ref_count += 1

    # Detect paper title from first # heading
    for line in lines:
        if line.startswith("# "):
            paper_title = line[2:].strip()
            break

    return {
        "paper_title": paper_title,
        "screenshots": screenshots,
        "figures": figures,
        "ref_count": ref_count,
        "md_path": str(md_path.resolve()),
        "paper_dir": str(md_path.parent.resolve()),
    }


# ── generator ────────────────────────────────────────────────

def generate_checklist(tasks: dict, output_path: Path) -> Path:
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title = tasks["paper_title"]
    screenshots = tasks["screenshots"]
    figures = tasks["figures"]

    # ── Title ──
    add_heading(doc, f"《{title}》— 待办清单", "title")

    # ── Section 1: Overview ──
    add_heading(doc, "一、当前状态", "heading1")
    done_count = len(figures)
    todo_count = len(screenshots)
    add_body(doc, f"论文已生成，当前共有 {done_count} 张实验图片已嵌入，{todo_count} 处代码截图待补充。以下列出所有需要你提供的文件和操作步骤。")

    # ── Section 2: Screenshot tasks ──
    if screenshots:
        add_heading(doc, "二、需要截图的代码段", "heading1")
        add_body(doc, f"请在北太天元中打开对应.m文件，截取代码区域，保存为PNG图片，放入指定路径。共 {len(screenshots)} 处：")

        for idx, ss in enumerate(screenshots, 1):
            # Figure number
            fig_num = ""
            for ch in ss.get("context", ""):
                if ch in "0123456789.-":
                    fig_num += ch
            add_heading(doc, f"任务 {idx}：{ss['label']}", "heading2")

            add_body(doc, f"截图内容：{ss['label']}")
            if ss.get("context"):
                add_body(doc, f"对应章节：{ss['context']}")
            add_body(doc, "操作步骤：")

            png_name = f"{ss['label']}.png"
            target_path = Path(tasks["paper_dir"]) / "figures" / png_name
            add_checkbox(doc, f"在北太天元中打开对应代码文件并截图", indent=True)
            add_checkbox(doc, f"截图保存为 PNG 格式", indent=True)
            add_checkbox(doc, f"将 PNG 文件放入：", indent=True)
            add_path(doc, f"    {target_path}")

            add_checkbox(doc, f"更新 image-map.json（或告诉我放入后我来更新）", indent=True)
            add_note(doc, f"截图完成后，将此文件路径告诉我：{target_path}")

    # ── Section 3: Image map update ──
    add_heading(doc, "三、需要更新图片映射", "heading1")
    add_body(doc, "截图放入 figures/ 目录后，运行以下命令更新图片映射：")
    add_path(doc, f"python scripts/build_image_map.py --scan {tasks['paper_dir']}/figures/ {tasks['paper_dir']}/image-map.json")
    add_body(doc, "或者直接告诉我图片已放好，我来重新生成 DOCX。")

    # ── Section 4: Final checklist ──
    add_heading(doc, "四、最终交付检查清单", "heading1")

    add_checkbox(doc, f"所有 {len(screenshots)} 处代码截图已放入 figures/ 目录")
    add_checkbox(doc, "image-map.json 已更新，包含所有新增截图")
    add_checkbox(doc, f"参考文献 {tasks['ref_count']} 条，全部在正文中有交叉引用")
    add_checkbox(doc, "DOCX 中无【待补】标记")
    add_checkbox(doc, f"全文字数在 3000-5000 字范围内")
    add_checkbox(doc, "格式符合学校要求（宋体/1.5倍行距/标题三号/正文小四）")
    add_checkbox(doc, "文件已提交给老师")

    # ── Section 5: File paths summary ──
    add_heading(doc, "五、关键文件路径速查", "heading1")

    paths = [
        ("论文 Markdown 源稿", f"{tasks['paper_dir']}/{tasks['paper_title']}.md"),
        ("论文 DOCX 输出", f"{tasks['paper_dir']}/{tasks['paper_title']}.docx"),
        ("图片目录", f"{tasks['paper_dir']}/figures/"),
        ("图片映射 JSON", f"{tasks['paper_dir']}/image-map.json"),
        ("论文元信息配置", f"{tasks['paper_dir']}/../templates/thesis-spec.yaml"),
        ("图表注册表", f"{tasks['paper_dir']}/../templates/figure-registry.yaml"),
        ("本项目 Skill", str(Path(__file__).resolve().parents[1])),
    ]
    for label, p in paths:
        add_body(doc, f"{label}：")
        add_path(doc, f"    {p}")

    # ── Section 6: What to tell me ──
    add_heading(doc, "六、完成后告诉我", "heading1")
    add_body(doc, "截图全部放好后，在对话中告诉我以下信息，我将立即重新生成完整版 DOCX：")

    items = [
        "截图已放入的目录路径（如：D:\\CourseWork\\人工智能导论\\paper-output\\figures\\）",
        "是否有需要替换或补充的图片",
        "是否需要调整论文章节内容或参考文献",
        "最终 DOCX 存放位置（默认同目录）",
    ]
    for item in items:
        add_checkbox(doc, item)

    # Save — write to temp first then copy (handles WSL cross-fs issues)
    import tempfile, shutil
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(tmp_path, str(output_path))
    finally:
        os.unlink(tmp_path)
    return output_path


# ── CLI ─────────────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser(description="Generate task checklist DOCX for course paper.")
    ap.add_argument("thesis_md", type=Path, help="Path to the thesis markdown file")
    ap.add_argument("--out", type=Path, default=None, help="Output DOCX path (default: Desktop/<title>-待办清单.docx)")
    args = ap.parse_args()

    md_path = args.thesis_md.resolve()
    if not md_path.exists():
        print(f"Error: {md_path} not found")
        return 1

    tasks = parse_tasks(md_path)

    if args.out:
        out = args.out.resolve()
    else:
        desktop = find_desktop()
        out = desktop / f"{tasks['paper_title']}-待办清单.docx"

    generate_checklist(tasks, out)

    print(f"✓ 待办清单已生成: {out}")
    print()
    print(f"  论文: {tasks['paper_title']}")
    print(f"  待补充截图: {len(tasks['screenshots'])} 处")
    print(f"  已嵌入图片: {len(tasks['figures'])} 张")
    print(f"  参考文献: {tasks['ref_count']} 条")
    print()
    print("  ── 接下来请你做 ──")
    for i, ss in enumerate(tasks['screenshots'], 1):
        print(f"  {i}. 截图「{ss['label']}」→ 放入 figures/ 目录")
    print()
    print(f"  完成后告诉我截图路径，我会重新生成完整 DOCX。")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
